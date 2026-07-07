import json
import shutil
import unittest
import zipfile
from pathlib import Path

from deep_book_learning.parsers import inspect_book, parse_book
from deep_book_learning.project import init_project
from deep_book_learning.validation import validate_project


class DeepBookLearningTests(unittest.TestCase):
    def setUp(self):
        self.tmp = Path.cwd() / ".test-tmp" / self.id().replace(".", "_")
        shutil.rmtree(self.tmp, ignore_errors=True)
        self.tmp.mkdir(parents=True, exist_ok=True)

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def write(self, name, text, encoding="utf-8"):
        p = self.tmp / name
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(text, encoding=encoding)
        return p

    def test_markdown_heading_extraction(self):
        p = self.write("sample.md", "# Root\nIntro\n\n## Chapter One\nMechanism text\n")
        parsed = parse_book(p)
        self.assertEqual(parsed.report.status, "PASS")
        self.assertGreaterEqual(parsed.report.chapter_count, 2)
        self.assertIn("Chapter One", [u.title for u in parsed.units])

    def test_txt_gb18030_encoding(self):
        p = self.tmp / "cn path.txt"
        p.write_bytes("\u7b2c\u4e00\u7ae0 \u5f00\u59cb\n\u8fd9\u662f\u6b63\u6587".encode("gb18030"))
        parsed = parse_book(p)
        self.assertIn(parsed.report.status, {"PASS", "WARNING"})
        self.assertIn("\u6b63\u6587", parsed.text)

    def test_empty_file_is_partial(self):
        p = self.write("empty.txt", "")
        report = inspect_book(p)
        self.assertEqual(report.status, "PARTIAL")

    def test_corrupt_epub_fails(self):
        p = self.tmp / "bad.epub"
        p.write_bytes(b"not a zip")
        report = inspect_book(p)
        self.assertEqual(report.status, "FAILED")

    def test_epub_spine_order(self):
        p = self.tmp / "book.epub"
        container = "<?xml version='1.0'?><container xmlns='urn:oasis:names:tc:opendocument:xmlns:container' version='1.0'><rootfiles><rootfile full-path='OEBPS/content.opf' media-type='application/oebps-package+xml'/></rootfiles></container>"
        opf = "<package xmlns='http://www.idpf.org/2007/opf' version='3.0'><metadata xmlns:dc='http://purl.org/dc/elements/1.1/'><dc:title>EPUB Fixture</dc:title><dc:creator>Tester</dc:creator><dc:language>en</dc:language></metadata><manifest><item id='c2' href='c2.xhtml' media-type='application/xhtml+xml'/><item id='c1' href='c1.xhtml' media-type='application/xhtml+xml'/></manifest><spine><itemref idref='c2'/><itemref idref='c1'/></spine></package>"
        with zipfile.ZipFile(p, "w") as z:
            z.writestr("META-INF/container.xml", container)
            z.writestr("OEBPS/content.opf", opf)
            z.writestr("OEBPS/c1.xhtml", "<html xmlns='http://www.w3.org/1999/xhtml'><body><h1>First file</h1><p>Should be second.</p></body></html>")
            z.writestr("OEBPS/c2.xhtml", "<html xmlns='http://www.w3.org/1999/xhtml'><body><h1>Second file</h1><p>Should be first.</p></body></html>")
        parsed = parse_book(p)
        self.assertEqual(parsed.report.status, "PASS")
        self.assertEqual([u.title for u in parsed.units], ["Second file", "First file"])

    def test_docx_heading_extraction_if_available(self):
        try:
            import docx
        except Exception:
            self.skipTest("python-docx unavailable")
        p = self.tmp / "sample.docx"
        d = docx.Document()
        d.add_heading("Chapter A", level=1)
        d.add_paragraph("Paragraph A")
        d.add_heading("Chapter B", level=1)
        d.add_paragraph("Paragraph B")
        d.save(str(p))
        parsed = parse_book(p)
        self.assertEqual(parsed.report.status, "PASS")
        self.assertIn("Chapter B", [u.title for u in parsed.units])

    def test_pdf_blank_detects_ocr_need_if_available(self):
        try:
            from pypdf import PdfWriter
        except Exception:
            self.skipTest("pypdf unavailable")
        p = self.tmp / "blank.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with p.open("wb") as f:
            writer.write(f)
        report = inspect_book(p)
        self.assertIn(report.status, {"WARNING", "PARTIAL"})
        self.assertTrue(report.warnings)

    def test_project_init_and_validation_with_chinese_path(self):
        cn_dir = self.tmp / "\u4e2d\u6587 \u8def\u5f84"
        cn_dir.mkdir()
        p = cn_dir / "sample book.md"
        p.write_text("# Book\nIntro\n\n## Model\nA mechanism explains outcomes.", encoding="utf-8")
        project = init_project(p, self.tmp / "out")
        ok, messages = validate_project(project)
        self.assertTrue(ok, messages)
        self.assertTrue((project / ".book_learning" / "manifest.json").exists())
        manifest = json.loads((project / ".book_learning" / "manifest.json").read_text(encoding="utf-8"))
        self.assertGreater(manifest["chunk_count"], 0)


if __name__ == "__main__":
    unittest.main()

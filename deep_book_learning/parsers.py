from __future__ import annotations

import hashlib
import html
import json
import mimetypes
import re
import zipfile
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from .models import BookUnit, InspectionReport, ParsedBook, SourceRef

SUPPORTED = {".pdf", ".epub", ".docx", ".txt", ".md", ".markdown", ".html", ".htm", ".rtf", ".odt"}


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def detect_format(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    mime = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    mapping = {
        ".pdf": "PDF",
        ".epub": "EPUB",
        ".docx": "DOCX",
        ".txt": "TXT",
        ".md": "Markdown",
        ".markdown": "Markdown",
        ".html": "HTML",
        ".htm": "HTML",
        ".rtf": "RTF",
        ".odt": "ODT",
    }
    return mapping.get(ext, "Unknown"), mime


def base_report(path: Path) -> InspectionReport:
    fmt, mime = detect_format(path)
    if not path.exists():
        return InspectionReport("FAILED", str(path), path.name, path.suffix.lower(), fmt, mime, 0, "", errors=["file does not exist"])
    return InspectionReport("WARNING", str(path), path.name, path.suffix.lower(), fmt, mime, path.stat().st_size, sha256_file(path))


def inspect_book(path_like: str | Path) -> InspectionReport:
    path = Path(path_like)
    report = base_report(path)
    if report.status == "FAILED":
        return report
    if report.extension not in SUPPORTED:
        report.status = "FAILED"
        report.errors.append(f"unsupported format: {report.extension or 'none'}")
        return report
    try:
        parsed = parse_book(path)
    except Exception as exc:  # keep CLI honest instead of crashing without context
        report.status = "FAILED"
        report.errors.append(f"parse failed: {exc.__class__.__name__}: {exc}")
        return report
    return parsed.report


def parse_book(path_like: str | Path) -> ParsedBook:
    path = Path(path_like)
    report = base_report(path)
    if report.extension not in SUPPORTED:
        report.status = "FAILED"
        report.errors.append(f"unsupported format: {report.extension or 'none'}")
        return ParsedBook(path, report, [])
    handlers = {
        ".pdf": _parse_pdf,
        ".epub": _parse_epub,
        ".docx": _parse_docx,
        ".txt": _parse_txt,
        ".md": _parse_markdown,
        ".markdown": _parse_markdown,
        ".html": _parse_html,
        ".htm": _parse_html,
        ".rtf": _parse_rtf,
        ".odt": _parse_odt,
    }
    return handlers[report.extension](path, report)


def _finish(report: InspectionReport, units: list[BookUnit]) -> ParsedBook:
    report.text_char_count = sum(len(u.text) for u in units)
    report.chapter_count = len(units)
    report.toc = [{"id": u.id, "title": u.title, "level": u.level, "char_count": len(u.text), "source": u.source.location} for u in units]
    if report.errors:
        report.status = "FAILED"
    elif report.text_char_count == 0:
        report.status = "PARTIAL"
        report.warnings.append("no extractable text found")
    elif report.warnings:
        report.status = "WARNING"
    else:
        report.status = "PASS"
    return ParsedBook(Path(report.path), report, units)


def _unit(i: int, title: str, text: str, label: str, location: dict, level: int = 2) -> BookUnit:
    unit_id = f"u{i:04d}"
    return BookUnit(unit_id, title.strip() or f"Unit {i}", level, text.strip(), SourceRef(unit_id, label, location))


def _parse_pdf(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "pypdf"
    try:
        from pypdf import PdfReader
    except Exception:
        report.status = "PARTIAL"
        report.errors.append("pypdf is not installed; cannot extract PDF text")
        return ParsedBook(path, report, [])
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        report.errors.append("PDF is encrypted")
        return _finish(report, [])
    meta = reader.metadata or {}
    report.title = str(meta.get("/Title") or path.stem) if meta else path.stem
    report.author = str(meta.get("/Author") or "") or None
    report.page_count = len(reader.pages)
    units = []
    blank_pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            text = ""
            report.warnings.append(f"page {i} extraction failed: {exc}")
        if not text.strip():
            blank_pages.append(i)
        units.append(_unit(i, f"Page {i}", text, f"page {i}", {"page": i}, level=5))
    if blank_pages:
        report.warnings.append(f"{len(blank_pages)} pages have no extractable text; OCR may be required")
        report.ocr_required = len(blank_pages) > max(1, len(reader.pages) // 3)
    report.source_strategy = "PDF page numbers"
    return _finish(report, units)


def _parse_docx(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "python-docx"
    try:
        import docx
    except Exception:
        report.errors.append("python-docx is not installed; cannot parse DOCX")
        return ParsedBook(path, report, [])
    doc = docx.Document(str(path))
    report.title = path.stem
    units = []
    current_title = "Front Matter"
    current = []
    idx = 1
    for pnum, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style else "") or ""
        if style.lower().startswith("heading") and current:
            units.append(_unit(idx, current_title, "\n\n".join(current), "DOCX paragraphs", {"paragraph_end": pnum - 1}))
            idx += 1
            current = []
            current_title = text
        elif style.lower().startswith("heading"):
            current_title = text
        else:
            current.append(text)
    if current or not units:
        units.append(_unit(idx, current_title, "\n\n".join(current), "DOCX paragraphs", {"paragraph_count": len(doc.paragraphs)}))
    report.source_strategy = "DOCX heading path and paragraph order"
    return _finish(report, units)


def _parse_txt(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "standard-library text"
    data = path.read_bytes()
    encodings = ["utf-8-sig", "utf-16", "gb18030", "big5", "latin-1"]
    text = None
    used = None
    for enc in encodings:
        try:
            text = data.decode(enc)
            used = enc
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        report.errors.append("could not decode text")
        return _finish(report, [])
    report.language = "unknown"
    report.source_strategy = f"line ranges ({used})"
    return _finish(report, _split_heading_text(text, path.stem, "line"))


def _parse_markdown(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "standard-library markdown"
    text = path.read_text(encoding="utf-8-sig")
    report.source_strategy = "Markdown heading path and line ranges"
    return _finish(report, _split_heading_text(text, path.stem, "line"))


def _parse_html(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "standard-library html-strip"
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = html.unescape(re.sub(r"<[^>]+>", " ", raw))
    report.warnings.append("HTML structure is simplified in v0.1.0")
    report.source_strategy = "HTML text order"
    return _finish(report, _split_heading_text(text, path.stem, "html"))


def _parse_rtf(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "basic rtf-strip"
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = re.sub(r"[{}]", " ", text)
    report.warnings.append("RTF parser is best-effort and may lose formatting")
    report.source_strategy = "RTF text order"
    return _finish(report, _split_heading_text(text, path.stem, "rtf"))


def _parse_odt(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "standard-library odt"
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("content.xml")
    text = _xml_text(xml)
    report.warnings.append("ODT parser is best-effort and flattens structure")
    report.source_strategy = "ODT content.xml order"
    return _finish(report, _split_heading_text(text, path.stem, "odt"))


def _parse_epub(path: Path, report: InspectionReport) -> ParsedBook:
    report.parser = "standard-library epub"
    units: list[BookUnit] = []
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        container = ET.fromstring(zf.read("META-INF/container.xml"))
        rootfile = container.find(".//{*}rootfile")
        if rootfile is None:
            report.errors.append("EPUB container has no rootfile")
            return _finish(report, [])
        opf_path = rootfile.attrib["full-path"]
        opf_dir = str(Path(opf_path).parent).replace(".", "")
        opf = ET.fromstring(zf.read(opf_path))
        report.title = _first_text(opf, ".//{*}title") or path.stem
        report.author = _first_text(opf, ".//{*}creator")
        report.language = _first_text(opf, ".//{*}language")
        manifest = {item.attrib.get("id"): item.attrib for item in opf.findall(".//{*}manifest/{*}item")}
        spine_ids = [item.attrib.get("idref") for item in opf.findall(".//{*}spine/{*}itemref")]
        idx = 1
        for idref in spine_ids:
            item = manifest.get(idref or "")
            if not item:
                continue
            href = item.get("href", "")
            media = item.get("media-type", "")
            if "html" not in media and not href.lower().endswith((".xhtml", ".html", ".htm")):
                continue
            full = str((Path(opf_dir) / href).as_posix()).lstrip("/") if opf_dir else href
            if full not in names:
                report.warnings.append(f"spine item missing: {full}")
                continue
            raw = zf.read(full)
            text = _xml_text(raw)
            title = _title_from_html(raw) or Path(href).stem
            units.append(_unit(idx, title, text, f"EPUB spine {idx}", {"spine_index": idx, "href": full}))
            idx += 1
    report.source_strategy = "EPUB spine order and href"
    if not units:
        report.errors.append("no readable EPUB spine XHTML found")
    return _finish(report, units)


def _first_text(root: ET.Element, path: str) -> str | None:
    el = root.find(path)
    return "".join(el.itertext()).strip() if el is not None else None


def _title_from_html(raw: bytes) -> str | None:
    try:
        root = ET.fromstring(raw)
        return _first_text(root, ".//{*}h1") or _first_text(root, ".//{*}title")
    except Exception:
        return None


def _xml_text(raw: bytes) -> str:
    try:
        root = ET.fromstring(raw)
        return html.unescape(" ".join(t.strip() for t in root.itertext() if t and t.strip()))
    except ET.ParseError:
        return html.unescape(re.sub(r"<[^>]+>", " ", raw.decode("utf-8", errors="replace")))


def _split_heading_text(text: str, default_title: str, source_key: str) -> list[BookUnit]:
    lines = text.splitlines()
    headings: list[tuple[int, str, int]] = []
    for i, line in enumerate(lines, start=1):
        stripped = line.strip()
        if re.match(r"^#{1,6}\s+", stripped):
            level = len(stripped) - len(stripped.lstrip("#"))
            headings.append((i, stripped.lstrip("# ").strip(), level))
        elif re.match(r"^(chapter|part|section)\s+\w+", stripped, flags=re.I):
            headings.append((i, stripped, 2))
        elif _looks_like_cjk_heading(stripped):
            headings.append((i, stripped, 2))
    if not headings:
        return [_unit(1, default_title, text, f"{source_key} 1-{len(lines)}", {"line_start": 1, "line_end": len(lines)}, 1)]
    units = []
    for idx, (start, title, level) in enumerate(headings, start=1):
        end = headings[idx][0] - 1 if idx < len(headings) else len(lines)
        body = "\n".join(lines[start - 1:end])
        units.append(_unit(idx, title, body, f"lines {start}-{end}", {"line_start": start, "line_end": end}, level))
    return units


def _looks_like_cjk_heading(value: str) -> bool:
    prefix = chr(0x7b2c)
    suffixes = {chr(0x7ae0), chr(0x8282), chr(0x90e8), chr(0x7bc7)}
    return bool(value.startswith(prefix) and any(ch in value[:16] for ch in suffixes))
